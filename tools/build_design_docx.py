import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = PROJECT_ROOT / "设计文档.md"


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=None, bold=None):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, first_line=True, alignment=None):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    if alignment is not None:
        paragraph.alignment = alignment


def clear_body_keep_section(doc):
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)


def configure_page(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    run = p.add_run("第 ")
    set_run_font(run, size=10.5)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)

    run2 = p.add_run(" 页")
    set_run_font(run2, size=10.5)


def parse_inline_code(text):
    parts = []
    pattern = re.compile(r"`([^`]+)`")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False))
        parts.append((m.group(1), True))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False))
    return parts or [(text, False)]


def add_text_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=True)
    for chunk, is_code in parse_inline_code(text):
        if not chunk:
            continue
        run = p.add_run(chunk)
        if is_code:
            set_run_font(run, east_asia="Consolas", ascii_font="Consolas", size=10.5)
            run.font.color.rgb = RGBColor(80, 80, 80)
        else:
            set_run_font(run, size=12)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.first_line_indent = None
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(12 if level == 1 else 6)
    fmt.space_after = Pt(6)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=16, bold=True)
    elif level == 2:
        set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=14, bold=True)
    else:
        set_run_font(run, east_asia="黑体", ascii_font="Times New Roman", size=12, bold=True)
    return p


def add_code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line.strip() else " ")
        set_run_font(run, east_asia="Consolas", ascii_font="Consolas", size=9.5)
        run.font.color.rgb = RGBColor(60, 60, 60)


def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.74)
    p.paragraph_format.first_line_indent = Cm(-0.32)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    marker = p.add_run("• ")
    set_run_font(marker, size=12)
    run = p.add_run(text)
    set_run_font(run, size=12)


def build_docx(template_path, output_path):
    doc = Document(str(template_path))
    clear_body_keep_section(doc)
    configure_page(doc)
    configure_styles(doc)

    md = SOURCE_MD.read_text(encoding="utf-8")
    lines = md.splitlines()

    in_code = False
    code_lines = []

    for raw in lines:
        line = raw.rstrip()
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), 1)
        elif line.startswith("## "):
            text = line[3:].strip()
            if re.match(r"第\d+章", text):
                doc.add_page_break()
                add_heading(doc, text, 1)
            else:
                add_heading(doc, text, 1)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 2)
        elif line.startswith("#### "):
            add_heading(doc, line[5:].strip(), 3)
        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
        else:
            add_text_paragraph(doc, line.strip())

    if code_lines:
        add_code_block(doc, code_lines)

    # Remove an accidental leading page break if the first content after title is a chapter.
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("}p"):
            p_pr = child.find(qn("w:pPr"))
            if p_pr is not None and p_pr.find(qn("w:pageBreakBefore")) is not None:
                p_pr.remove(p_pr.find(qn("w:pageBreakBefore")))
            break

    add_page_number(doc.sections[0])
    doc.save(str(output_path))


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_design_docx.py <template.docx> <output.docx>")
    build_docx(Path(sys.argv[1]), Path(sys.argv[2]))
