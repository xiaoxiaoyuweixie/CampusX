import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_run_font(run, east_asia="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run.font.size = Pt(size)
    run.bold = bold


def clear_page_break_before(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:pageBreakBefore"))
    if node is not None:
        p_pr.remove(node)


def set_page_break_before(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:pageBreakBefore")) is None:
        node = OxmlElement("w:pageBreakBefore")
        p_pr.append(node)


def normalize_runs(paragraph, east_asia, ascii_font, size, bold=False):
    if not paragraph.runs:
        paragraph.add_run(paragraph.text)
    for run in paragraph.runs:
        set_run_font(run, east_asia=east_asia, ascii_font=ascii_font, size=size, bold=bold)


def is_code_line(text):
    stripped = text.strip()
    if not stripped:
        return False
    return (
        stripped in {"{", "}", "},"}
        or stripped.startswith('"')
        or stripped.startswith(("├", "│", "└"))
        or stripped.startswith("CampusX")
        or "=> callFunction" in stripped
    )


def finalize(path):
    doc = Document(str(path))
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.4)

    for index, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        fmt = p.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = 1.5

        if index == 0 and text.startswith("CampusX"):
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.first_line_indent = None
            fmt.space_after = Pt(12)
            normalize_runs(p, "黑体", "Times New Roman", 18, True)
            continue

        if re.match(r"^第\d+章", text) or text == "参考资料":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.first_line_indent = None
            fmt.space_before = Pt(12)
            fmt.space_after = Pt(6)
            if index > 1 and text != "参考资料":
                set_page_break_before(p)
            normalize_runs(p, "黑体", "Times New Roman", 16, True)
            continue

        if re.match(r"^\d+\.\d+\s", text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = None
            fmt.space_before = Pt(6)
            fmt.space_after = Pt(3)
            clear_page_break_before(p)
            normalize_runs(p, "黑体", "Times New Roman", 14, True)
            continue

        if re.match(r"^\d+\.\d+\.\d+\s", text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = None
            fmt.space_before = Pt(4)
            fmt.space_after = Pt(2)
            clear_page_break_before(p)
            normalize_runs(p, "黑体", "Times New Roman", 12, True)
            continue

        if is_code_line(text):
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = None
            fmt.left_indent = Cm(0.74)
            fmt.line_spacing = 1.15
            normalize_runs(p, "Consolas", "Consolas", 9.5, False)
            continue

        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.first_line_indent = Cm(0.74)
        normalize_runs(p, "宋体", "Times New Roman", 12, False)

    doc.save(str(path))


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("Usage: finalize_design_docx.py <docx>")
    finalize(Path(sys.argv[1]))
